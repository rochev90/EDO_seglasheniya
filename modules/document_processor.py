#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os, re, sys
from datetime import datetime
from typing import Dict
from docx import Document


class DocumentProcessor:
    """Создание документов из docx-шаблонов с подстановкой данных."""

    def __init__(self, output_folder: str = "Соглашения"):
        # Шаблоны всегда в <директория_программы>/data
        app_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
        self.templates_folder = os.path.join(app_dir, "data")
        self.output_folder = output_folder
        self.months = {
            1: "января", 2: "февраля", 3: "марта", 4: "апреля",
            5: "мая", 6: "июня", 7: "июля", 8: "августа",
            9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"
        }
        self.templates = {
            ("КАДИС", "IP"): "KADIS_IP_shablon.docx",
            ("КАДИС", "OOO"): "KADIS_OOO_shablon.docx",
            ("ЮрРегионИнформ", "IP"): "URI_IP_shablon.docx",
            ("ЮрРегионИнформ", "OOO"): "URI_OOO_shablon.docx",
        }

    @staticmethod
    def _safe_filename(text: str) -> str:
        name = re.sub(r'[\\/\:\*\?\"\<\>\|\r\n\t]+', ' ', str(text))
        name = re.sub(r'\s+', ' ', name).strip(' .')
        return name[:140]

    def _ensure_output_dir(self, company: str) -> str:
        date_folder = datetime.now().strftime("%d.%m.%y")
        path = os.path.join(self.output_folder, company, date_folder)
        os.makedirs(path, exist_ok=True)
        return path

    def _date_mapping(self) -> Dict[str, str]:
        now = datetime.now()
        return {
            "{{dd}}": f"{now.day:02d}",
            "{{mm}}": self.months[now.month],
            "{{yy}}": f"{now.year}",
            "{dd}": f"{now.day:02d}",
            "{mm}": self.months[now.month],
            "{yy}": f"{now.year}",
        }

    @staticmethod
    def format_fio_short(full_fio: str) -> str:
        """
        Преобразует ФИО в формат И.О. Фамилия

        Примеры:
            "Иванов Иван Иванович" -> "И.И. Иванов"
            "ИП Иванов Иван Иванович" -> "И.И. Иванов"
            "Петров Петр" -> "П. Петров"
            "Гулиев Парвиз Октай Оглы" -> "П.О. Гулиев"
        """
        # Убираем "ИП" если есть
        fio = full_fio.strip()
        if fio.upper().startswith("ИП "):
            fio = fio[3:].strip()

        # Разделяем на части
        parts = fio.split()

        if len(parts) == 0:
            return fio

        # Фамилия всегда первая
        surname = parts[0]

        # Игнорируемые окончания (если это отдельное слово)
        ignore_endings = {"оглы", "кызы", "углы", "кизы"}

        # Инициалы из имени и отчества
        initials = []
        for part in parts[1:]:
            if part and len(part) > 0:
                # Проверяем, не является ли это игнорируемым окончанием
                if part.lower() not in ignore_endings:
                    initials.append(part[0].upper() + ".")

        # Формируем результат
        if initials:
            # Инициалы без пробелов между ними, затем пробел и фамилия
            return "".join(initials) + " " + surname
        else:
            return surname

    @staticmethod
    def _replace_in_paragraph(paragraph, mapping: Dict[str, str]):
        if not paragraph.text:
            return
        text = paragraph.text
        for k, v in mapping.items():
            text = text.replace(k, v)
        if text != paragraph.text:
            for r in paragraph.runs:
                r.text = ''
            paragraph.runs[0].text = text

    def _apply_mapping(self, doc: Document, mapping: Dict[str, str]):
        for p in doc.paragraphs:
            self._replace_in_paragraph(p, mapping)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_paragraph(p, mapping)

    def _open_template(self, company: str, kind: str) -> Document:
        fname = self.templates[(company, kind)]
        path = os.path.join(self.templates_folder, fname)
        if not os.path.exists(path):
            raise FileNotFoundError(f"Шаблон не найден: {fname} (ищу только в: {self.templates_folder})")
        return Document(path)

    def fill_ip_template(self, company: str, ip_name: str, ip_inn: str, fio: str) -> str:
        """
        Заполняет шаблон для ИП.

        Args:
            company: название компании (КАДИС, ЮрРегионИнформ)
            ip_name: полное название ИП (например "ИП Иванов Иван Иванович")
            ip_inn: ИНН
            fio: ФИО (будет преобразовано в И.О. Фамилия)
        """
        doc = self._open_template(company, "IP")

        # Форматируем ФИО в И.О. Фамилия
        fio_short = self.format_fio_short(fio)

        mapping = {
            "{{IP}}": ip_name,
            "{{IP_INN}}": ip_inn,
            "{{fio}}": fio_short
        }
        mapping.update(self._date_mapping())
        self._apply_mapping(doc, mapping)

        out_dir = self._ensure_output_dir(company)
        out_name = f"Соглашение_ЭДО_{self._safe_filename(ip_name)}.docx"
        out_path = os.path.join(out_dir, out_name)
        doc.save(out_path)
        return out_path

    def fill_ul_template(self, company: str, org_name: str, inn: str, kpp: str,
                         position: str, fio: str, post_fixed: str, fio_fixed: str) -> str:
        """
        Заполняет шаблон для юридического лица.

        Args:
            company: название компании (КАДИС, ЮрРегионИнформ)
            org_name: название организации
            inn: ИНН
            kpp: КПП
            position: должность (именительный падеж)
            fio: ФИО (именительный падеж, будет преобразовано в И.О. Фамилия)
            post_fixed: должность (родительный падеж)
            fio_fixed: ФИО (родительный падеж, полное)
        """
        doc = self._open_template(company, "OOO")

        # Форматируем только fio в И.О. Фамилия
        # fio_fixed остается полным ФИО в родительном падеже
        fio_short = self.format_fio_short(fio)

        mapping = {
            "{{JL}}": org_name,
            "{{JL_INN}}": inn,
            "{{JL_KPP}}": kpp,
            "{{post}}": position,
            "{{fio}}": fio_short,
            "{{post_fixed}}": post_fixed,
            "{{fio_fixed}}": fio_fixed,
        }
        mapping.update(self._date_mapping())
        self._apply_mapping(doc, mapping)

        out_dir = self._ensure_output_dir(company)
        out_name = f"Соглашение_ЭДО_{self._safe_filename(org_name)}.docx"
        out_path = os.path.join(out_dir, out_name)
        doc.save(out_path)
        return out_path