#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Скрипт для создания примеров шаблонов документов
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_ip_template(filename, company_name):
    """Создает шаблон для ИП"""
    doc = Document()
    
    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"СОГЛАШЕНИЕ\nоб электронном документообороте")
    run.bold = True
    run.font.size = Pt(14)
    
    # Дата и место
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"г. Санкт-Петербург                                                {{dd}} {{{{mm}}}} {{{{yy}}}} г.")
    
    doc.add_paragraph()
    
    # Основной текст
    p1 = doc.add_paragraph()
    p1.add_run(f"{company_name}, ИНН 7710324108, в лице генерального директора Иванова Ивана Ивановича, "
               f"действующего на основании Устава, именуемое в дальнейшем «Заказчик», "
               f"с одной стороны, и {{{{IP}}}}, ИНН {{{{IP_INN}}}}, именуемый в дальнейшем «Исполнитель», "
               f"в лице {{{{fio}}}}, действующего без доверенности, "
               f"с другой стороны, совместно именуемые «Стороны», заключили настоящее Соглашение о нижеследующем:")
    
    doc.add_paragraph()
    
    # Пункты соглашения
    doc.add_heading("1. ПРЕДМЕТ СОГЛАШЕНИЯ", level=2)
    
    p2 = doc.add_paragraph()
    p2.add_run("1.1. Стороны договорились об организации электронного документооборота между собой "
               "в рамках выполнения договорных обязательств.")
    
    p3 = doc.add_paragraph()
    p3.add_run("1.2. Электронный документооборот осуществляется через оператора электронного документооборота "
               "АО «ПФ «СКБ Контур» с использованием системы «Контур.Диадок».")
    
    doc.add_heading("2. РЕКВИЗИТЫ СТОРОН", level=2)
    
    # Таблица реквизитов
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # Заголовки
    table.rows[0].cells[0].text = "Заказчик:"
    table.rows[0].cells[1].text = "Исполнитель:"
    
    # Реквизиты
    table.rows[1].cells[0].text = f"{company_name}\nИНН: 7710324108\nКПП: 781301001"
    table.rows[1].cells[1].text = "{{IP}}\nИНН: {{IP_INN}}"
    
    doc.add_paragraph()
    
    # Подписи
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.rows[0].cells[0].text = "Заказчик:"
    sig_table.rows[0].cells[1].text = "Исполнитель:"
    sig_table.rows[1].cells[0].text = "Генеральный директор"
    sig_table.rows[1].cells[1].text = "{{fio}}"
    sig_table.rows[2].cells[0].text = "___________ / Иванов И.И. /"
    sig_table.rows[2].cells[1].text = "___________ / {{fio}} /"
    
    # Сохранение
    doc.save(filename)
    print(f"Создан шаблон: {filename}")

def create_ul_template(filename, company_name):
    """Создает шаблон для ЮЛ"""
    doc = Document()
    
    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"СОГЛАШЕНИЕ\nоб электронном документообороте")
    run.bold = True
    run.font.size = Pt(14)
    
    # Дата и место
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"г. Санкт-Петербург                                                {{dd}} {{{{mm}}}} {{{{yy}}}} г.")
    
    doc.add_paragraph()
    
    # Основной текст
    p1 = doc.add_paragraph()
    p1.add_run(f"{company_name}, ИНН 7710324108, в лице генерального директора Иванова Ивана Ивановича, "
               f"действующего на основании Устава, именуемое в дальнейшем «Заказчик», "
               f"с одной стороны, и {{{{JL}}}}, ИНН {{{{JL_INN}}}}, КПП {{{{JL_KPP}}}}, "
               f"в лице {{{{post_fixed}}}} {{{{fio_fixed}}}}, действующего на основании Устава, "
               f"именуемое в дальнейшем «Исполнитель», "
               f"с другой стороны, совместно именуемые «Стороны», заключили настоящее Соглашение о нижеследующем:")
    
    doc.add_paragraph()
    
    # Пункты соглашения
    doc.add_heading("1. ПРЕДМЕТ СОГЛАШЕНИЯ", level=2)
    
    p2 = doc.add_paragraph()
    p2.add_run("1.1. Стороны договорились об организации электронного документооборота между собой "
               "в рамках выполнения договорных обязательств.")
    
    p3 = doc.add_paragraph()
    p3.add_run("1.2. Электронный документооборот осуществляется через оператора электронного документооборота "
               "АО «ПФ «СКБ Контур» с использованием системы «Контур.Диадок».")
    
    doc.add_heading("2. РЕКВИЗИТЫ СТОРОН", level=2)
    
    # Таблица реквизитов
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # Заголовки
    table.rows[0].cells[0].text = "Заказчик:"
    table.rows[0].cells[1].text = "Исполнитель:"
    
    # Реквизиты
    table.rows[1].cells[0].text = f"{company_name}\nИНН: 7710324108\nКПП: 781301001"
    table.rows[1].cells[1].text = "{{JL}}\nИНН: {{JL_INN}}\nКПП: {{JL_KPP}}"
    
    doc.add_paragraph()
    
    # Подписи
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.rows[0].cells[0].text = "Заказчик:"
    sig_table.rows[0].cells[1].text = "Исполнитель:"
    sig_table.rows[1].cells[0].text = "Генеральный директор"
    sig_table.rows[1].cells[1].text = "{{post}}"
    sig_table.rows[2].cells[0].text = "___________ / Иванов И.И. /"
    sig_table.rows[2].cells[1].text = "___________ / {{fio}} /"
    
    # Сохранение
    doc.save(filename)
    print(f"Создан шаблон: {filename}")

if __name__ == "__main__":
    # Создаем папку data если её нет
    os.makedirs("data", exist_ok=True)
    
    print("Создание примеров шаблонов...")
    print()
    
    # Шаблоны для КАДИС
    create_ip_template("data/KADIS_IP_shablon.docx", "ООО \"КАДИС\"")
    create_ul_template("data/KADIS_OOO_shablon.docx", "ООО \"КАДИС\"")
    
    # Шаблоны для ЮрРегионИнформ
    create_ip_template("data/URI_IP_shablon.docx", "ООО \"ЮрРегионИнформ\"")
    create_ul_template("data/URI_OOO_shablon.docx", "ООО \"ЮрРегионИнформ\"")
    
    print()
    print("✓ Все шаблоны успешно созданы!")
    print("Примечание: Это примеры шаблонов. Замените их на реальные шаблоны вашей компании.")
